import json
import sqlite3
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from config import (
    FIXED_JSON_DIR,
    SUMMARY_DIR,
    STATUS_DOCX_DONE,
    get_fixed_json_dir,
    get_db_path,
    page_url_to_site_name,
    safe_name,
)


# =========================================================
# 1. 基础工具
# =========================================================
def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def flatten_json(obj):
    """
    递归展开 JSON，提取所有 dict
    """
    results = []

    if isinstance(obj, dict):
        results.append(obj)
        for v in obj.values():
            results.extend(flatten_json(v))
    elif isinstance(obj, list):
        for item in obj:
            results.extend(flatten_json(item))

    return results


def crawl_time_str(crawl_time: datetime) -> str:
    return crawl_time.strftime("%Y-%m-%d %H:%M:%S")


def crawl_time_file_str(crawl_time: datetime) -> str:
    return crawl_time.strftime("%Y%m%d_%H%M%S")


def build_log_docx_name(page_url: str, crawl_time: datetime) -> str:
    """
    文档名：爬取的网址 + 爬取时间
    """
    name = safe_name(f"{page_url}_{crawl_time_file_str(crawl_time)}")
    return f"{name[:180]}.docx"


# =========================================================
# 2. 固定 JSON 里的记录筛选
# =========================================================
def is_valid_content_record(d: dict) -> bool:
    """
    判断一个 dict 是否是有效内容条目
    """
    if not isinstance(d, dict):
        return False

    item_id = d.get("itemId") or d.get("item_id")
    title = d.get("title")
    url = d.get("url")
    item_type = str(d.get("itemType", "")).strip()
    data_valid = d.get("dataValid", True)

    if not item_id:
        return False

    if not isinstance(title, str) or not title.strip():
        return False

    if not isinstance(url, str) or not url.strip():
        return False

    if "xuexi.cn/lgpage/detail/index.html" not in url:
        return False

    if item_type.lower() == "outerlink":
        return False

    if data_valid is False:
        return False

    return True


def is_video_record(d: dict) -> bool:
    """
    判断是否是视频内容
    """
    item_type = str(d.get("itemType", "")).strip()
    content_type = str(d.get("type", "")).strip()

    return content_type == "shipin" or item_type == "kPureVideo"


def extract_video_records(data: dict) -> list[dict]:
    """
    从固定 JSON 中提取视频记录
    """
    records = []
    dicts = flatten_json(data)

    for d in dicts:
        if not is_valid_content_record(d):
            continue

        if not is_video_record(d):
            continue

        item_id = str(d.get("itemId") or d.get("item_id"))
        title = str(d.get("title", "")).strip()
        publish_time = str(d.get("publishTime", "")).strip()

        records.append({
            "item_id": item_id,
            "title": title,
            "publish_time": publish_time,
        })

    unique = {}
    for r in records:
        unique[r["item_id"]] = r

    return list(unique.values())


# =========================================================
# 3. 读取 db 状态
# =========================================================
def load_video_state_from_db(db_path: Path) -> dict[str, dict]:
    """
    从运行数据库读取每个 item_id 的处理状态和生成文档地址。
    """
    if not db_path.exists():
        return {}

    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row

    try:
        rows = conn.execute("""
            SELECT item_id, status, docx_path
            FROM videos
        """).fetchall()
    except sqlite3.Error:
        rows = []
    finally:
        conn.close()

    state_map = {}
    for row in rows:
        state_map[row["item_id"]] = {
            "status": row["status"],
            "docx_path": row["docx_path"] or "",
        }

    return state_map


# =========================================================
# 4. 生成爬取日志数据
# =========================================================
def build_crawl_log_rows_for_site(
    site_name: str,
    target_item_ids: set[str] | None = None,
) -> list[dict]:
    """
    合并同一个网址下本次处理的视频记录。
    target_item_ids 为 None 时保留兼容逻辑，读取全部视频记录。
    """
    fixed_json_dir = get_fixed_json_dir(site_name)
    if not fixed_json_dir.exists():
        raise FileNotFoundError(f"固定 JSON 目录不存在：{fixed_json_dir}")

    rows_by_item_id = {}

    for json_path in sorted(fixed_json_dir.glob("*.json")):
        json_name = json_path.name
        db_path = get_db_path(site_name, json_name)
        state_map = load_video_state_from_db(db_path)

        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        for record in extract_video_records(data):
            item_id = record["item_id"]
            if target_item_ids is not None and item_id not in target_item_ids:
                continue

            state = state_map.get(item_id, {})
            raw_status = state.get("status", "NEW")
            is_done = raw_status == STATUS_DOCX_DONE

            row = {
                "item_id": item_id,
                "title": record["title"],
                "publish_time": record["publish_time"],
                "is_done": is_done,
                "status": "已完成" if is_done else "未完成",
                "docx_path": state.get("docx_path", ""),
            }

            existing = rows_by_item_id.get(item_id)
            if existing is None or (not existing["is_done"] and row["is_done"]):
                rows_by_item_id[item_id] = row

    def sort_key(row: dict) -> str:
        return row.get("publish_time") or ""

    return sorted(rows_by_item_id.values(), key=sort_key, reverse=True)


def save_crawl_log_docx(page_url: str, rows: list[dict], crawl_time: datetime) -> Path:
    """
    在爬取日志目录下直接生成 docx，不再按网站名创建子目录。
    """
    ensure_dir(SUMMARY_DIR)

    log_path = SUMMARY_DIR / build_log_docx_name(page_url, crawl_time)
    done_count = sum(1 for row in rows if row["is_done"])
    pending_count = len(rows) - done_count

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("爬取日志")
    title_run.bold = True
    title_run.font.size = Pt(16)

    doc.add_paragraph(f"爬取网址：{page_url}")
    doc.add_paragraph(f"爬取时间：{crawl_time_str(crawl_time)}")
    doc.add_paragraph(f"视频总数：{len(rows)}，已完成：{done_count}，未完成：{pending_count}")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = True

    headers = table.rows[0].cells
    headers[0].text = "视频标题"
    headers[1].text = "是否完成爬取"
    headers[2].text = "文档存储地址"

    for cell in headers:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row_data in rows:
        cells = table.add_row().cells
        cells[0].text = row_data["title"]
        cells[1].text = row_data["status"]
        cells[2].text = row_data["docx_path"] or "未生成"

    doc.save(log_path)
    print(f"[爬取日志] 已保存：{log_path}")
    return log_path


# =========================================================
# 5. 刷新爬取日志
# =========================================================
def discover_summary_targets(
    target_page_urls: list[str] | None = None,
    target_site_names: list[str] | None = None,
) -> list[tuple[str, str]]:
    """
    返回要刷新爬取日志的目标：
    [(site_name, page_url), ...]
    """
    if target_page_urls:
        return [(page_url_to_site_name(url), url) for url in target_page_urls]

    if target_site_names:
        return [(site_name, site_name) for site_name in target_site_names]

    targets = []
    if not FIXED_JSON_DIR.exists():
        return targets

    for site_dir in FIXED_JSON_DIR.iterdir():
        if site_dir.is_dir():
            targets.append((site_dir.name, site_dir.name))

    return targets


def refresh_summaries(
    target_page_urls: list[str] | None = None,
    target_site_names: list[str] | None = None,
    processed_item_ids_by_site: dict[str, list[str]] | None = None,
) -> list[dict]:
    """
    刷新爬取日志：
    - 传入 target_page_urls 时，用真实网址生成文件名和文档内容
    - 传入 processed_item_ids_by_site 时，只记录本轮实际处理的视频
    - 只传 target_site_names 时，用网站目录名作为网址兜底
    - 都不传时，刷新全部网站
    """
    crawl_time = datetime.now()
    targets = discover_summary_targets(
        target_page_urls=target_page_urls,
        target_site_names=target_site_names,
    )

    if not targets:
        print("没有找到可刷新爬取日志的目标。")
        return []

    results = []

    for site_name, page_url in targets:
        try:
            target_item_ids = None
            if processed_item_ids_by_site is not None:
                target_item_ids = set(processed_item_ids_by_site.get(site_name, []))

            rows = build_crawl_log_rows_for_site(
                site_name=site_name,
                target_item_ids=target_item_ids,
            )
            log_path = save_crawl_log_docx(page_url, rows, crawl_time)
            results.append({
                "site_name": site_name,
                "page_url": page_url,
                "crawl_log_path": str(log_path),
                "video_count": len(rows),
                "done_count": sum(1 for row in rows if row["is_done"]),
                "pending_count": sum(1 for row in rows if not row["is_done"]),
                "status": "SUCCESS",
            })
        except Exception as e:
            results.append({
                "site_name": site_name,
                "page_url": page_url,
                "crawl_log_path": None,
                "status": "FAILED",
                "error": str(e),
            })
            print(f"[爬取日志 FAILED] {site_name} -> {e}")

    return results
