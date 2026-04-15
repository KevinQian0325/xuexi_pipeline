import base64
import copy
import json
import re
import sqlite3
import subprocess
import time
import uuid
from pathlib import Path
import shutil

import requests
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from pydub import AudioSegment
from pydub.silence import detect_silence

from config import (
    DB_DIR,
    APP_ID,
    ACCESS_TOKEN,
    RECOGNIZE_URL,
    RESOURCE_ID,
    AUDIO_SAMPLE_RATE,
    AUDIO_CHANNELS,
    AUDIO_CODEC,
    MAX_AUDIO_DURATION_MS,
    MAX_AUDIO_FILE_SIZE,
    CHUNK_TARGET_MS,
    CUT_SEARCH_BACK_MS,
    MIN_SILENCE_LEN_MS,
    KEEP_SILENCE_MS,
    SILENCE_THRESH_OFFSET_DB,
    STATUS_M3U8_DONE,
    STATUS_VIDEO_DONE,
    STATUS_AUDIO_DONE,
    STATUS_ASR_DONE,
    STATUS_DOCX_DONE,
    STATUS_FAILED,
    page_url_to_site_name,
    get_video_material_dir,
)
from capture_m3u8 import get_first_m3u8
from build_index import create_table


STEP_M3U8 = "M3U8_CAPTURE"
STEP_VIDEO = "VIDEO_DOWNLOAD"
STEP_AUDIO = "AUDIO_EXTRACT"
STEP_ASR = "ASR_RECOGNIZE"
STEP_DOCX = "DOCX_GENERATE"


# =========================================================
# 1. 基础工具
# =========================================================
def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)

def remove_dir_if_exists(path: Path) -> None:
    if path.exists() and path.is_dir():
        shutil.rmtree(path)

def now_str() -> str:
    return time.strftime("%Y-%m-%d %H:%M:%S")


def safe_name(text: str, max_len: int = 80) -> str:
    text = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", str(text))
    text = re.sub(r"_+", "_", text).strip("_")
    if not text:
        text = "video"
    return text[:max_len]


def file_to_base64(file_path: Path) -> str:
    with open(file_path, "rb") as f:
        file_data = f.read()
    return base64.b64encode(file_data).decode("utf-8")


def format_ms(ms: int) -> str:
    total_seconds = ms // 1000
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    seconds = total_seconds % 60
    if hours > 0:
        return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
    return f"{minutes:02d}:{seconds:02d}"


def split_long_text(text: str, max_len: int = 110) -> list[str]:
    parts = re.split(r"(?<=[。！？])", text)
    result = []
    buf = ""

    for part in parts:
        if len(buf) + len(part) <= max_len:
            buf += part
        else:
            if buf.strip():
                result.append(buf.strip())
            buf = part

    if buf.strip():
        result.append(buf.strip())

    return result


def get_audio_duration_ms_and_size(audio_path: Path) -> tuple[int, int]:
    audio = AudioSegment.from_wav(audio_path)
    return len(audio), audio_path.stat().st_size


def existing_file(path: Path) -> bool:
    return path.exists() and path.is_file() and path.stat().st_size > 0


def get_existing_output_path(row: sqlite3.Row, field_name: str, default_path: Path) -> Path:
    stored_path = row[field_name]
    if stored_path:
        candidate = Path(stored_path)
        if existing_file(candidate):
            return candidate

    return default_path


# =========================================================
# 2. DB 操作
# =========================================================
def update_video_record(conn: sqlite3.Connection, item_id: str, **fields) -> None:
    """
    动态更新 videos 表中的字段
    """
    if not fields:
        return

    fields["updated_at"] = now_str()

    set_clause = ", ".join([f"{k} = ?" for k in fields.keys()])
    params = list(fields.values()) + [item_id]

    sql = f"UPDATE videos SET {set_clause} WHERE item_id = ?"
    conn.execute(sql, params)
    conn.commit()


def mark_video_attempt(
    conn: sqlite3.Connection,
    item_id: str,
    run_id: str,
    page_url: str,
    json_name: str,
    material_dir: Path,
) -> None:
    """
    标记某条视频进入本轮处理。
    """
    conn.execute("""
    UPDATE videos
    SET
        attempt_count = COALESCE(attempt_count, 0) + 1,
        last_run_id = ?,
        last_processed_at = ?,
        source_page_url = ?,
        source_json_name = ?,
        material_dir = ?,
        error_step = NULL,
        error_type = NULL,
        error_message = NULL,
        updated_at = ?
    WHERE item_id = ?
    """, (
        run_id,
        now_str(),
        page_url,
        json_name,
        str(material_dir),
        now_str(),
        item_id,
    ))
    conn.commit()


def record_video_event(
    conn: sqlite3.Connection,
    run_id: str,
    item_id: str,
    step: str,
    status: str,
    message: str | None = None,
) -> None:
    """
    记录单条视频的阶段事件，便于后期排查。
    """
    conn.execute("""
    INSERT INTO video_events (
        run_id,
        item_id,
        step,
        status,
        message,
        created_at
    )
    VALUES (?, ?, ?, ?, ?, ?)
    """, (
        run_id,
        item_id,
        step,
        status,
        message,
        now_str(),
    ))
    conn.commit()


def create_or_update_crawl_run(
    conn: sqlite3.Connection,
    run_id: str,
    page_url: str,
    process_start_time: str | None,
    process_end_time: str | None,
    started_at: str,
    target_count: int,
) -> None:
    conn.execute("""
    INSERT INTO crawl_runs (
        run_id,
        page_url,
        process_start_time,
        process_end_time,
        started_at,
        status,
        target_count,
        success_count,
        failed_count
    )
    VALUES (?, ?, ?, ?, ?, ?, ?, 0, 0)
    ON CONFLICT(run_id) DO UPDATE SET
        page_url = excluded.page_url,
        process_start_time = excluded.process_start_time,
        process_end_time = excluded.process_end_time,
        started_at = excluded.started_at,
        status = excluded.status,
        target_count = excluded.target_count
    """, (
        run_id,
        page_url,
        process_start_time,
        process_end_time,
        started_at,
        "RUNNING",
        target_count,
    ))
    conn.commit()


def finish_crawl_run(
    conn: sqlite3.Connection,
    run_id: str,
    success_count: int,
    failed_count: int,
    error_message: str | None = None,
) -> None:
    if failed_count > 0 and success_count > 0:
        status = "PARTIAL_FAILED"
    elif failed_count > 0:
        status = "FAILED"
    else:
        status = "SUCCESS"

    conn.execute("""
    UPDATE crawl_runs
    SET
        finished_at = ?,
        status = ?,
        success_count = ?,
        failed_count = ?,
        error_message = ?
    WHERE run_id = ?
    """, (
        now_str(),
        status,
        success_count,
        failed_count,
        error_message,
        run_id,
    ))
    conn.commit()


def update_crawl_log_path(db_path: Path, run_id: str, crawl_log_path: str) -> None:
    conn = sqlite3.connect(db_path)
    create_table(conn)
    conn.execute("""
    UPDATE crawl_runs
    SET crawl_log_path = ?
    WHERE run_id = ?
    """, (
        crawl_log_path,
        run_id,
    ))
    conn.commit()
    conn.close()


def load_pending_videos(
    conn: sqlite3.Connection,
    start_time: str | None = None,
    end_time: str | None = None,
) -> list[sqlite3.Row]:
    """
    读取未完成视频。
    如果没给时间范围，就返回全部未完成视频。
    当前唯一支持的筛选条件就是 publish_time 时间范围。
    """
    sql = """
    SELECT *
    FROM videos
    WHERE status != ?
    """
    params = [STATUS_DOCX_DONE]

    if start_time:
        sql += " AND publish_time >= ?"
        params.append(start_time)

    if end_time:
        sql += " AND publish_time <= ?"
        params.append(end_time)

    sql += " ORDER BY publish_time DESC"

    conn.row_factory = sqlite3.Row
    rows = conn.execute(sql, params).fetchall()
    return rows


def discover_db_targets(target_page_urls: list[str] | None = None) -> list[tuple[str, str, Path]]:
    """
    返回要处理的 db 列表：
    [(site_name, json_name, db_path), ...]

    规则：
    - target_page_urls is None -> 处理运行数据库目录下全部网站的全部 db
    - target_page_urls 有值 -> 只处理这些网站名对应目录下的全部 db
    """
    targets = []

    if not DB_DIR.exists():
        return targets

    target_site_names = None
    if target_page_urls:
        target_site_names = {page_url_to_site_name(url) for url in target_page_urls}

    for site_dir in DB_DIR.iterdir():
        if not site_dir.is_dir():
            continue

        site_name = site_dir.name

        if target_site_names is not None and site_name not in target_site_names:
            continue

        for db_path in site_dir.glob("*.db"):
            json_name = db_path.stem + ".json"
            targets.append((site_name, json_name, db_path))

    return targets


# =========================================================
# 3. ffmpeg
# =========================================================
def run_ffmpeg(cmd: list[str]) -> None:
    subprocess.run(cmd, check=True)


def m3u8_to_mp4(m3u8_url: str, mp4_path: Path) -> None:
    cmd = [
        "ffmpeg",
        "-y",
        "-protocol_whitelist", "file,http,https,tcp,tls,crypto",
        "-i", m3u8_url,
        "-c", "copy",
        str(mp4_path),
    ]
    print(f"[FFMPEG] 开始下载视频 -> {mp4_path.name}")
    run_ffmpeg(cmd)
    print(f"[FFMPEG] 视频下载完成: {mp4_path}")


def mp4_to_wav(mp4_path: Path, wav_path: Path) -> None:
    cmd = [
        "ffmpeg",
        "-y",
        "-i", str(mp4_path),
        "-vn",
        "-acodec", AUDIO_CODEC,
        "-ar", str(AUDIO_SAMPLE_RATE),
        "-ac", str(AUDIO_CHANNELS),
        str(wav_path),
    ]
    print(f"[FFMPEG] 开始从视频提取音频 -> {wav_path.name}")
    run_ffmpeg(cmd)
    print(f"[FFMPEG] 音频提取完成: {wav_path}")


# =========================================================
# 4. ASR
# =========================================================
def recognize_flash(file_path: Path) -> dict:
    if not APP_ID or not ACCESS_TOKEN:
        raise ValueError("请先在 .env 中配置 XUEXI_APP_ID 和 XUEXI_ACCESS_TOKEN。")

    headers = {
        "X-Api-App-Key": APP_ID,
        "X-Api-Access-Key": ACCESS_TOKEN,
        "X-Api-Resource-Id": RESOURCE_ID,
        "X-Api-Request-Id": str(uuid.uuid4()),
        "X-Api-Sequence": "-1",
    }

    audio_data = {
        "data": file_to_base64(file_path)
    }

    request_body = {
        "user": {
            "uid": APP_ID
        },
        "audio": audio_data,
        "request": {
            "model_name": "bigmodel",
            "enable_itn": True,
            "enable_punc": True,
            "enable_ddc": True,
            "show_utterances": True,
        },
    }

    print(f"[ASR] 开始识别: {file_path.name}")
    response = requests.post(
        RECOGNIZE_URL,
        json=request_body,
        headers=headers,
        timeout=600
    )

    if "X-Api-Status-Code" not in response.headers:
        raise RuntimeError(f"接口未返回标准响应头，headers={response.headers}")

    code = response.headers["X-Api-Status-Code"]
    msg = response.headers.get("X-Api-Message", "")
    logid = response.headers.get("X-Tt-Logid", "")

    print(f"[ASR] code={code}, message={msg}, logid={logid}")

    if code != "20000000":
        try:
            body = response.json()
        except Exception:
            body = response.text
        raise RuntimeError(f"识别失败: code={code}, msg={msg}, logid={logid}, body={body}")

    return response.json()


def build_chunk_ranges(audio: AudioSegment) -> list[tuple[int, int]]:
    """
    规则：
    1. 以 30 分钟为一个目标段长
    2. 到达目标点后，从目标点往前找最近的静音点
    3. 找到静音点就在静音点切
    4. 找不到就按目标点硬切
    5. 最后一段如果不足30分钟，直接输出
    """
    total_ms = len(audio)
    silence_thresh = audio.dBFS - SILENCE_THRESH_OFFSET_DB

    ranges = []
    start_ms = 0

    while start_ms < total_ms:
        target_end = min(start_ms + CHUNK_TARGET_MS, total_ms)

        if target_end >= total_ms:
            ranges.append((start_ms, total_ms))
            break

        search_start = max(start_ms, target_end - CUT_SEARCH_BACK_MS)
        search_segment = audio[search_start:target_end]

        silence_ranges = detect_silence(
            search_segment,
            min_silence_len=MIN_SILENCE_LEN_MS,
            silence_thresh=silence_thresh,
        )

        if silence_ranges:
            last_silence_start, _last_silence_end = silence_ranges[-1]
            cut_point = search_start + last_silence_start
            if cut_point <= start_ms:
                cut_point = target_end
        else:
            cut_point = target_end

        ranges.append((start_ms, cut_point))
        start_ms = cut_point

    return ranges


def export_audio_chunks(wav_path: Path, chunk_dir: Path) -> list[tuple[Path, int]]:
    """
    返回：
    [
        (chunk_path, chunk_start_offset_ms),
        ...
    ]
    """
    ensure_dir(chunk_dir)
    audio = AudioSegment.from_wav(wav_path)
    ranges = build_chunk_ranges(audio)

    exported = []
    for idx, (start_ms, end_ms) in enumerate(ranges, start=1):
        chunk_start = max(0, start_ms - KEEP_SILENCE_MS)
        chunk_end = min(len(audio), end_ms + KEEP_SILENCE_MS)

        chunk = audio[chunk_start:chunk_end]
        chunk_path = chunk_dir / f"part_{idx:03d}.wav"
        chunk.export(chunk_path, format="wav")
        exported.append((chunk_path, chunk_start))

        print(
            f"[CHUNK] 导出 {chunk_path.name} | "
            f"start={chunk_start}ms end={chunk_end}ms "
            f"duration={(chunk_end - chunk_start) / 1000:.2f}s"
        )

    return exported


def merge_asr_results(results_with_offsets: list[tuple[dict, int]], total_duration_ms: int) -> dict:
    """
    把多段 ASR 结果合并成一个和原接口基本一致的 JSON
    """
    merged_text_parts = []
    merged_utterances = []

    for result, offset_ms in results_with_offsets:
        text = result.get("result", {}).get("text", "")
        if text:
            merged_text_parts.append(text)

        for utt in result.get("result", {}).get("utterances", []) or []:
            new_utt = copy.deepcopy(utt)

            if "start_time" in new_utt:
                new_utt["start_time"] += offset_ms
            if "end_time" in new_utt:
                new_utt["end_time"] += offset_ms

            for word in new_utt.get("words", []) or []:
                if "start_time" in word:
                    word["start_time"] += offset_ms
                if "end_time" in word:
                    word["end_time"] += offset_ms

            merged_utterances.append(new_utt)

    merged_text = "".join(merged_text_parts)

    return {
        "audio_info": {
            "duration": total_duration_ms
        },
        "result": {
            "additions": {
                "duration": str(total_duration_ms)
            },
            "text": merged_text,
            "utterances": merged_utterances,
        }
    }


def recognize_flash_smart(wav_path: Path, chunk_dir: Path) -> dict:
    """
    不超限就直接识别；
    超限就切块后逐段识别，再合并结果。
    """
    duration_ms, file_size = get_audio_duration_ms_and_size(wav_path)

    print(f"[CHECK] 音频时长: {duration_ms / 1000:.2f} 秒")
    print(f"[CHECK] 音频大小: {file_size / 1024 / 1024:.2f} MB")

    if duration_ms <= MAX_AUDIO_DURATION_MS and file_size <= MAX_AUDIO_FILE_SIZE:
        print("[ASR] 未超限，直接识别整段音频")
        return recognize_flash(wav_path)

    print("[ASR] 音频超限，开始切块识别")
    chunk_files = export_audio_chunks(wav_path, chunk_dir)

    results_with_offsets = []
    for chunk_path, offset_ms in chunk_files:
        chunk_duration_ms, chunk_size = get_audio_duration_ms_and_size(chunk_path)

        print(
            f"[ASR-CHUNK] {chunk_path.name} | "
            f"duration={chunk_duration_ms / 1000:.2f}s | "
            f"size={chunk_size / 1024 / 1024:.2f}MB"
        )

        if chunk_duration_ms > MAX_AUDIO_DURATION_MS or chunk_size > MAX_AUDIO_FILE_SIZE:
            raise RuntimeError(
                f"切块后仍然超限: {chunk_path.name}, "
                f"duration={chunk_duration_ms / 1000:.2f}s, "
                f"size={chunk_size / 1024 / 1024:.2f}MB"
            )

        chunk_result = recognize_flash(chunk_path)
        results_with_offsets.append((chunk_result, offset_ms))

    merged = merge_asr_results(results_with_offsets, total_duration_ms=duration_ms)
    print("[ASR] 切块识别完成，已合并结果")
    return merged


# =========================================================
# 5. Word 生成
# =========================================================
def set_doc_font(document: Document, western: str = "Arial", east_asia: str = "Microsoft YaHei") -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = western
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    normal_style.font.size = Pt(10.5)


def add_custom_style(document: Document, style_name: str, size: float, bold: bool, color: str) -> None:
    if style_name in document.styles:
        style = document.styles[style_name]
    else:
        style = document.styles.add_style(style_name, 1)

    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def json_to_word(data: dict, output_docx: Path, title: str) -> None:
    duration_ms = data["audio_info"]["duration"]
    full_text = data["result"]["text"]
    utterances = data["result"].get("utterances", [])

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    set_doc_font(doc)
    add_custom_style(doc, "MyTitle", 18, True, "1F4E79")
    add_custom_style(doc, "MyHeading1", 13, True, "1F4E79")
    add_custom_style(doc, "MyMeta", 9, False, "666666")

    p = doc.add_paragraph(style="MyTitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)

    p = doc.add_paragraph(style="MyMeta")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"音频时长：{format_ms(duration_ms)}（约 {duration_ms / 1000:.1f} 秒）")

    p = doc.add_paragraph(style="MyHeading1")
    p.add_run("一、全文整理")

    for para_text in split_long_text(full_text):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(para_text)

    p = doc.add_paragraph(style="MyHeading1")
    p.add_run("二、分段时间轴")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False

    table.columns[0].width = Cm(2.6)
    table.columns[1].width = Cm(2.6)
    table.columns[2].width = Cm(11.2)

    headers = table.rows[0].cells
    headers[0].text = "开始"
    headers[1].text = "结束"
    headers[2].text = "内容"

    for cell in headers:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
        shade_cell(cell, "D9EAF7")

    for utt in utterances:
        row = table.add_row().cells
        row[0].text = format_ms(utt["start_time"])
        row[1].text = format_ms(utt["end_time"])
        row[2].text = utt["text"]

        for i, cell in enumerate(row):
            for p in cell.paragraphs:
                if i < 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = "Arial"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    r.font.size = Pt(9.5)

    p = doc.add_paragraph(style="MyHeading1")
    p.add_run("三、说明")

    notes = [
        "本稿依据识别接口返回结果自动整理，便于阅读、复制和归档。",
        "如需进一步整理为纯正文稿、逐段发言稿或字幕稿，可继续在此基础上处理。"
    ]

    for note in notes:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.add_run(note)

    doc.save(output_docx)


# =========================================================
# 6. 单条视频处理
# =========================================================
def process_one_video(
    conn: sqlite3.Connection,
    site_name: str,
    json_name: str,
    row: sqlite3.Row,
    run_id: str,
    page_url: str,
) -> str:
    item_id = row["item_id"]
    title = row["title"]
    publish_time = row["publish_time"]
    detail_url = row["detail_url"]

    material_dir = get_video_material_dir(site_name, json_name, title, publish_time)
    ensure_dir(material_dir)

    mp4_path = material_dir / "视频.mp4"
    wav_path = material_dir / "音频.wav"
    docx_path = material_dir / "文本.docx"
    chunk_dir = material_dir / "_chunks"

    mp4_path = get_existing_output_path(row, "mp4_path", mp4_path)
    wav_path = get_existing_output_path(row, "wav_path", wav_path)
    docx_path = get_existing_output_path(row, "docx_path", docx_path)

    print("=" * 100)
    print(f"[VIDEO] 开始处理 item_id={item_id}")
    print(f"标题: {title}")
    print(f"发布时间: {publish_time}")

    mark_video_attempt(
        conn=conn,
        item_id=item_id,
        run_id=run_id,
        page_url=page_url,
        json_name=json_name,
        material_dir=material_dir,
    )

    current_step = "INIT"

    try:
        if existing_file(docx_path):
            update_video_record(
                conn,
                item_id,
                docx_path=str(docx_path),
                status=STATUS_DOCX_DONE,
                docx_done_at=now_str(),
                error_message=None,
            )
            record_video_event(conn, run_id, item_id, STEP_DOCX, "SKIPPED", f"复用已有文档：{docx_path}")
            print(f"[SKIP] 已存在 Word 文档，直接标记完成: {docx_path}")
            return STATUS_DOCX_DONE

        has_existing_wav = existing_file(wav_path)
        has_existing_mp4 = existing_file(mp4_path)

        # 1) 抓 m3u8
        m3u8_url = row["m3u8_url"]
        if m3u8_url or has_existing_mp4 or has_existing_wav:
            update_fields = {
                "status": STATUS_M3U8_DONE,
                "error_message": None,
            }
            if m3u8_url:
                update_fields["m3u8_url"] = m3u8_url
                skip_message = "复用数据库中的 m3u8_url"
            else:
                skip_message = "已有本地媒体文件，无需重新捕获 m3u8"

            update_video_record(conn, item_id, **update_fields)
            record_video_event(conn, run_id, item_id, STEP_M3U8, "SKIPPED", skip_message)
            print(f"[SKIP] {skip_message}")
        else:
            current_step = STEP_M3U8
            record_video_event(conn, run_id, item_id, current_step, "START")
            m3u8_url = get_first_m3u8(detail_url)
            if not m3u8_url:
                raise RuntimeError("未抓到 m3u8 地址")

            update_video_record(
                conn,
                item_id,
                m3u8_url=m3u8_url,
                status=STATUS_M3U8_DONE,
                m3u8_done_at=now_str(),
                error_message=None,
            )
            record_video_event(conn, run_id, item_id, current_step, "SUCCESS", m3u8_url)

        # 2) 下载 mp4
        if has_existing_wav:
            record_video_event(conn, run_id, item_id, STEP_VIDEO, "SKIPPED", f"已有音频，跳过视频下载：{wav_path}")
            print(f"[SKIP] 已存在 wav，跳过视频下载: {wav_path}")
        elif has_existing_mp4:
            update_video_record(
                conn,
                item_id,
                mp4_path=str(mp4_path),
                mp4_size=mp4_path.stat().st_size,
                status=STATUS_VIDEO_DONE,
                error_message=None,
            )
            record_video_event(conn, run_id, item_id, STEP_VIDEO, "SKIPPED", f"复用已有视频：{mp4_path}")
            print(f"[SKIP] 已存在 mp4，跳过下载: {mp4_path}")
        else:
            current_step = STEP_VIDEO
            record_video_event(conn, run_id, item_id, current_step, "START")
            m3u8_to_mp4(m3u8_url, mp4_path)
            update_video_record(
                conn,
                item_id,
                mp4_path=str(mp4_path),
                mp4_size=mp4_path.stat().st_size if mp4_path.exists() else None,
                status=STATUS_VIDEO_DONE,
                video_done_at=now_str(),
            )
            record_video_event(conn, run_id, item_id, current_step, "SUCCESS", str(mp4_path))

        # 3) 提取 wav
        if existing_file(wav_path):
            audio_duration_ms, wav_size = get_audio_duration_ms_and_size(wav_path)
            update_video_record(
                conn,
                item_id,
                wav_path=str(wav_path),
                wav_size=wav_size,
                audio_duration_ms=audio_duration_ms,
                status=STATUS_AUDIO_DONE,
                error_message=None,
            )
            record_video_event(conn, run_id, item_id, STEP_AUDIO, "SKIPPED", f"复用已有音频：{wav_path}")
            print(f"[SKIP] 已存在 wav，跳过音频提取: {wav_path}")
        else:
            current_step = STEP_AUDIO
            record_video_event(conn, run_id, item_id, current_step, "START")
            mp4_to_wav(mp4_path, wav_path)
            audio_duration_ms, wav_size = get_audio_duration_ms_and_size(wav_path)
            update_video_record(
                conn,
                item_id,
                wav_path=str(wav_path),
                wav_size=wav_size,
                audio_duration_ms=audio_duration_ms,
                status=STATUS_AUDIO_DONE,
                audio_done_at=now_str(),
            )
            record_video_event(conn, run_id, item_id, current_step, "SUCCESS", str(wav_path))

        # 4) ASR（整段或切块）
        current_step = STEP_ASR
        record_video_event(conn, run_id, item_id, current_step, "START")
        asr_result = recognize_flash_smart(wav_path, chunk_dir)
        update_video_record(
            conn,
            item_id,
            status=STATUS_ASR_DONE,
            asr_done_at=now_str(),
        )
        record_video_event(conn, run_id, item_id, current_step, "SUCCESS")

        # 5) 生成 Word
        current_step = STEP_DOCX
        record_video_event(conn, run_id, item_id, current_step, "START")
        json_to_word(asr_result, docx_path, title=title)
        update_video_record(
            conn,
            item_id,
            docx_path=str(docx_path),
            status=STATUS_DOCX_DONE,
            docx_done_at=now_str(),
            error_message=None,
        )
        record_video_event(conn, run_id, item_id, current_step, "SUCCESS", str(docx_path))

        # 成功后删除切块目录
        remove_dir_if_exists(chunk_dir)

        print(f"[DONE] item_id={item_id} 处理完成")
        return STATUS_DOCX_DONE

    except Exception as e:
        # 失败也删除切块目录
        remove_dir_if_exists(chunk_dir)

        error_message = str(e)
        update_video_record(
            conn,
            item_id,
            status=STATUS_FAILED,
            error_step=current_step,
            error_type=type(e).__name__,
            error_message=error_message,
        )
        record_video_event(conn, run_id, item_id, current_step, "FAILED", error_message)
        print(f"[FAILED] item_id={item_id} -> {e}")
        return STATUS_FAILED


# =========================================================
# 7. 处理一个 db
# =========================================================
def process_one_db(
    site_name: str,
    json_name: str,
    db_path: Path,
    start_time: str | None = None,
    end_time: str | None = None,
    run_id: str | None = None,
    page_url: str | None = None,
    run_started_at: str | None = None,
) -> dict:
    print("\n" + "#" * 100)
    print(f"开始处理 DB：{db_path}")
    print(f"时间范围：start_time={start_time}, end_time={end_time}")

    run_id = run_id or str(uuid.uuid4())
    page_url = page_url or site_name
    run_started_at = run_started_at or now_str()

    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    create_table(conn)

    pending_rows = load_pending_videos(
        conn,
        start_time=start_time,
        end_time=end_time,
    )

    print(f"待处理视频数：{len(pending_rows)}")

    processed_item_ids = [row["item_id"] for row in pending_rows]
    create_or_update_crawl_run(
        conn=conn,
        run_id=run_id,
        page_url=page_url,
        process_start_time=start_time,
        process_end_time=end_time,
        started_at=run_started_at,
        target_count=len(pending_rows),
    )

    final_statuses = []
    for row in pending_rows:
        final_status = process_one_video(
            conn=conn,
            site_name=site_name,
            json_name=json_name,
            row=row,
            run_id=run_id,
            page_url=page_url,
        )
        final_statuses.append(final_status)

    run_success_count = sum(1 for status in final_statuses if status == STATUS_DOCX_DONE)
    run_failed_count = sum(1 for status in final_statuses if status == STATUS_FAILED)
    finish_crawl_run(
        conn=conn,
        run_id=run_id,
        success_count=run_success_count,
        failed_count=run_failed_count,
    )

    remain = conn.execute(
        "SELECT COUNT(*) FROM videos WHERE status != ?",
        (STATUS_DOCX_DONE,)
    ).fetchone()[0]

    done = conn.execute(
        "SELECT COUNT(*) FROM videos WHERE status = ?",
        (STATUS_DOCX_DONE,)
    ).fetchone()[0]

    failed = conn.execute(
        "SELECT COUNT(*) FROM videos WHERE status = ?",
        (STATUS_FAILED,)
    ).fetchone()[0]

    conn.close()

    return {
        "site_name": site_name,
        "json_name": json_name,
        "db_path": str(db_path),
        "run_id": run_id,
        "page_url": page_url,
        "processed_count": len(pending_rows),
        "run_success_count": run_success_count,
        "run_failed_count": run_failed_count,
        "done_count": done,
        "failed_count": failed,
        "remaining_count": remain,
        "processed_item_ids": processed_item_ids,
    }


# =========================================================
# 8. 以网站为单位处理
# =========================================================
def process_sites(
    target_page_urls: list[str] | None = None,
    start_time: str | None = None,
    end_time: str | None = None,
    run_id: str | None = None,
    run_started_at: str | None = None,
) -> list[dict]:
    """
    由 pipeline 调用：
    - target_page_urls = None -> 处理全部网站
    - target_page_urls 有值 -> 只处理这些网站对应目录下的全部 db

    时间范围是当前唯一支持的搜索条件：
    - 不给时间范围 -> 默认处理全部未完成视频
    - 给 start/end -> 按 publish_time 过滤
    """
    targets = discover_db_targets(target_page_urls=target_page_urls)

    if not targets:
        print("没有找到可处理的 db 目标。")
        return []

    results = []
    run_id = run_id or str(uuid.uuid4())
    run_started_at = run_started_at or now_str()

    page_url_by_site = {}
    if target_page_urls:
        page_url_by_site = {
            page_url_to_site_name(url): url
            for url in target_page_urls
        }

    for site_name, json_name, db_path in targets:
        result = process_one_db(
            site_name=site_name,
            json_name=json_name,
            db_path=db_path,
            start_time=start_time,
            end_time=end_time,
            run_id=run_id,
            page_url=page_url_by_site.get(site_name, site_name),
            run_started_at=run_started_at,
        )
        results.append(result)

    return results
